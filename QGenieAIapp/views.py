import os
import json
import secrets
import uuid
from django.shortcuts import render, redirect, get_object_or_404
from django.contrib.auth import login, authenticate, logout, update_session_auth_hash
from django.contrib.auth.models import User
from django.contrib.auth.decorators import login_required
from django.contrib import messages
from django.db.models import Avg, Count, Max, Sum
from django.http import Http404
from django.core.mail import send_mail
from django.utils import timezone
from django.core.files.base import ContentFile

from .models import UploadedFile, Question, NormalQuestion, QuizResult, OTPVerification, UserProfile, BookmarkedQuestion
from .forms import UserRegisterForm, UserLoginForm, UploadedFileForm, OTPVerificationForm, UserUpdateForm, UserProfileUpdateForm
from .ai_service import extract_text_from_pdf, generate_mcqs, generate_normal_questions, chat_with_document, generate_summary

# --- User Authentication Views ---

def register_view(request):
    """
    Registers a new user (with is_active=False), generates a 6-digit OTP code,
    sends it to their email, and redirects them to the verification page.
    """
    if request.user.is_authenticated:
        return redirect('dashboard')
        
    if request.method == 'POST':
        form = UserRegisterForm(request.POST)
        if form.is_valid():
            # Save user with is_active = False so they must verify their email first
            user = form.save(commit=False)
            user.is_active = False
            user.save()
            
            # Generate a secure 6-digit OTP
            otp_code = str(secrets.randbelow(900000) + 100000)
            expires_at = timezone.now() + timezone.timedelta(minutes=10)
            
            # Save OTP to database
            OTPVerification.objects.update_or_create(
                user=user,
                defaults={'otp_code': otp_code, 'expires_at': expires_at}
            )
            
            # Send Email
            subject = "Verify Your QGenie.AI Account - OTP Code"
            message = f"Hello {user.username},\n\nThank you for registering at QGenie.AI!\n\nYour 6-digit OTP verification code is: {otp_code}\n\nThis code will expire in 10 minutes.\n\nBest regards,\nQGenie.AI Team"
            
            try:
                send_mail(
                    subject,
                    message,
                    None,  # Uses DEFAULT_FROM_EMAIL from settings
                    [user.email],
                    fail_silently=False,
                )
                messages.success(request, "Registration successful! A 6-digit verification code has been sent to your email.")
            except Exception as e:
                messages.warning(request, f"Registration successful, but we could not send the email: {str(e)}. Please click 'Resend OTP' on the verification page.")
                
            return redirect('verify_otp', username=user.username)
        else:
            for field, errors in form.errors.items():
                for error in errors:
                    messages.error(request, f"{field.capitalize()}: {error}")
    else:
        form = UserRegisterForm()
        
    return render(request, 'user_panel/register.html', {'form': form})


def verify_otp_view(request, username):
    """
    Handles OTP validation. Activates the user and logs them in upon successful validation.
    """
    user = get_object_or_404(User, username=username)
    
    if not user.is_active and not hasattr(user, 'otp_verification'):
        messages.error(request, "This account has been suspended/blocked by the administrator.")
        return redirect('login')
        
    if user.is_active:
        messages.info(request, "Your account is already active. Please log in.")
        return redirect('login')
        
    if request.method == 'POST':
        form = OTPVerificationForm(request.POST)
        if form.is_valid():
            otp_code = form.cleaned_data.get('otp_code')
            
            try:
                verification = user.otp_verification
                if verification.otp_code == otp_code:
                    if verification.is_expired():
                        messages.error(request, "This OTP code has expired. Please request a new one.")
                    else:
                        # Success! Activate user
                        user.is_active = True
                        user.save()
                        
                        # Clean up OTP record
                        verification.delete()
                        
                        # Log user in
                        login(request, user)
                        messages.success(request, f"Account verified successfully! Welcome to QGenie.AI, {user.username}!")
                        return redirect('dashboard')
                else:
                    messages.error(request, "Invalid OTP code. Please try again.")
            except OTPVerification.DoesNotExist:
                messages.error(request, "No OTP verification details found for this user. Please request a new code.")
        else:
            for field, errors in form.errors.items():
                for error in errors:
                    messages.error(request, f"{field.capitalize()}: {error}")
    else:
        form = OTPVerificationForm()
        
    return render(request, 'user_panel/verify_otp.html', {'form': form, 'username': username})


def resend_otp_view(request, username):
    """
    Generates a new OTP, updates the DB, and resends the verification email.
    """
    user = get_object_or_404(User, username=username)
    
    if not user.is_active and not hasattr(user, 'otp_verification'):
        messages.error(request, "This account has been suspended/blocked by the administrator.")
        return redirect('login')
        
    if user.is_active:
        messages.info(request, "Your account is already active. Please log in.")
        return redirect('login')
        
    otp_code = str(secrets.randbelow(900000) + 100000)
    expires_at = timezone.now() + timezone.timedelta(minutes=10)
    
    OTPVerification.objects.update_or_create(
        user=user,
        defaults={'otp_code': otp_code, 'expires_at': expires_at}
    )
    
    subject = "Your New QGenie.AI OTP Code"
    message = f"Hello {user.username},\n\nHere is your new 6-digit OTP verification code: {otp_code}\n\nThis code will expire in 10 minutes.\n\nBest regards,\nQGenie.AI Team"
    
    try:
        send_mail(
            subject,
            message,
            None,
            [user.email],
            fail_silently=False,
        )
        messages.success(request, "A new 6-digit OTP code has been sent to your email.")
    except Exception as e:
        messages.error(request, f"Failed to send verification email: {str(e)}")
        
    return redirect('verify_otp', username=username)



def login_view(request):
    """
    Logs in an existing user and redirects them to the dashboard or admin dashboard based on their role.
    """
    if request.user.is_authenticated:
        if request.user.is_staff:
            return redirect('admin_dashboard')
        return redirect('dashboard')
        
    if request.method == 'POST':
        form = UserLoginForm(request, data=request.POST)
        if form.is_valid():
            username = form.cleaned_data.get('username')
            password = form.cleaned_data.get('password')
            user = authenticate(username=username, password=password)
            if user is not None:
                login(request, user)
                messages.success(request, f"Welcome back, {username}!")
                if user.is_staff:
                    return redirect('admin_dashboard')
                return redirect('dashboard')
        else:
            # Check if login failed because user is inactive (unverified or suspended)
            username = request.POST.get('username', '').strip()
            try:
                user = User.objects.get(username=username)
                if not user.is_active:
                    # Check if they have an active OTP record (unverified vs suspended)
                    if hasattr(user, 'otp_verification'):
                        messages.warning(request, "Your account has not been verified yet. Please verify your email below.")
                        return redirect('verify_otp', username=user.username)
                    else:
                        messages.error(request, "Your account has been suspended/blocked by the administrator.")
                        return redirect('login')
            except User.DoesNotExist:
                pass
                
            messages.error(request, "Invalid username or password.")
    else:
        form = UserLoginForm()
        
    return render(request, 'user_panel/login.html', {'form': form})


def logout_view(request):
    """
    Logs out the user and redirects them to the login page.
    """
    logout(request)
    messages.info(request, "You have been logged out successfully.")
    return redirect('login')


# --- Core Web App User Views ---

@login_required
def dashboard_view(request):
    """
    Renders the central user dashboard showing:
    - User stats (total uploads, quizzes completed, average score)
    - List of uploaded files with quiz options
    - List of recent quiz results (history)
    """
    user = request.user
    
    # User uploads and results
    uploads = UploadedFile.objects.filter(user=user).order_by('-uploaded_at')
    results = QuizResult.objects.filter(user=user).order_by('-date_taken')
    
    # Calculate key performance stats
    total_uploads = uploads.count()
    quizzes_completed = results.count()
    
    avg_score = results.aggregate(Avg('percentage'))['percentage__avg'] or 0.0
    avg_score = round(avg_score, 1)

    # Advanced Analytics for Interactive Progress Charts
    chart_results = list(reversed(results.order_by('-date_taken')[:15]))
    chart_data = {
        'labels': [res.date_taken.strftime('%b %d') for res in chart_results],
        'percentages': [res.percentage for res in chart_results],
        'titles': [res.file.title if res.file else 'Deleted File' for res in chart_results]
    }
    chart_data_json = json.dumps(chart_data)

    # Performance breakdown by difficulty level
    easy_avg = results.filter(file__difficulty='easy').aggregate(Avg('percentage'))['percentage__avg'] or 0.0
    medium_avg = results.filter(file__difficulty='medium').aggregate(Avg('percentage'))['percentage__avg'] or 0.0
    hard_avg = results.filter(file__difficulty='hard').aggregate(Avg('percentage'))['percentage__avg'] or 0.0
    
    easy_avg = round(easy_avg, 1)
    medium_avg = round(medium_avg, 1)
    hard_avg = round(hard_avg, 1)

    # Helper stats
    high_score = results.aggregate(Max('percentage'))['percentage__max'] or 0.0
    high_score = round(high_score, 1)
    
    passing_quizzes = results.filter(percentage__gte=60.0).count()
    passing_rate = round((passing_quizzes / quizzes_completed) * 100, 1) if quizzes_completed > 0 else 0.0
    
    total_questions_answered = results.aggregate(Sum('total_questions'))['total_questions__sum'] or 0

    context = {
        'uploads': uploads,
        'results': results[:5],  # Show recent 5 attempts on dashboard
        'total_uploads': total_uploads,
        'quizzes_completed': quizzes_completed,
        'avg_score': avg_score,
        'chart_data_json': chart_data_json,
        'easy_avg': easy_avg,
        'medium_avg': medium_avg,
        'hard_avg': hard_avg,
        'high_score': high_score,
        'passing_rate': passing_rate,
        'total_questions_answered': total_questions_answered,
    }
    return render(request, 'user_panel/dashboard.html', context)


@login_required
def upload_view(request):
    """
    Handles PDF and TXT file uploads OR pasted text input, and generates
    a custom number of MCQs using Gemini AI.
    """
    if request.method == 'POST':
        # Get standard upload inputs
        title = request.POST.get('title')
        input_type = request.POST.get('input_type', 'file')
        difficulty = request.POST.get('difficulty', 'medium')
        if difficulty not in ['easy', 'medium', 'hard']:
            difficulty = 'medium'
            
        try:
            num_questions = int(request.POST.get('num_questions', 5))
            gen_type = request.POST.get('gen_type', 'mcq')  # 'mcq', 'question', 'both'
            include_answers = request.POST.get('include_answers', 'yes') == 'yes'
            after_generate = request.POST.get('after_generate', 'review')  # 'review' or 'quiz'
            if num_questions not in [5, 10, 15, 20]:
                num_questions = 5
        except ValueError:
            num_questions = 5
            gen_type = 'mcq'
            include_answers = True
            after_generate = 'review'

        instance = None
        extracted_text = ""

        try:
            if input_type == 'text':
                pasted_text = request.POST.get('pasted_text', '').strip()
                pasted_title = request.POST.get('pasted_title', '').strip()

                if len(pasted_text) < 30:
                    messages.error(request, "Pasted text must contain at least 30 characters.")
                    return redirect('upload')

                title = pasted_title or f"Pasted Text ({timezone.now().strftime('%d %b %Y, %H:%M')})"
                
                # Create the model instance first
                instance = UploadedFile(
                    user=request.user,
                    title=title,
                    file_type='TXT',
                    include_answers=include_answers,
                    difficulty=difficulty
                )
                
                # Create a virtual text file and save it
                filename = f"pasted_{uuid.uuid4().hex[:10]}.txt"
                instance.file.save(filename, ContentFile(pasted_text.encode('utf-8')), save=False)
                instance.save()
                
                extracted_text = pasted_text

            else:  # input_type == 'file'
                form = UploadedFileForm(request.POST, request.FILES)
                if not form.is_valid():
                    messages.error(request, "Please select a valid PDF or TXT file.")
                    return redirect('upload')

                uploaded_file_obj = request.FILES['file']
                filename = uploaded_file_obj.name
                
                # File validation
                if not filename.lower().endswith(('.pdf', '.txt')):
                    messages.error(request, "Invalid file type. Only PDF and TXT files are supported.")
                    return redirect('upload')
                    
                # Create instance of UploadedFile (without saving yet to extract metadata)
                instance = form.save(commit=False)
                instance.user = request.user
                instance.title = os.path.splitext(filename)[0]
                instance.file_type = 'PDF' if filename.lower().endswith('.pdf') else 'TXT'
                instance.include_answers = include_answers
                instance.difficulty = difficulty
                instance.save()
                
                # Text extraction logic
                if instance.file_type == 'PDF':
                    extracted_text = extract_text_from_pdf(instance.file.path)
                else:
                    # TXT file processing
                    with open(instance.file.path, 'r', encoding='utf-8', errors='ignore') as f:
                        extracted_text = f.read()

                # Basic validation on content length
                if len(extracted_text.strip()) < 30:
                    instance.delete()  # Clean up DB
                    messages.error(request, "The uploaded file contains insufficient readable text. Please try another.")
                    return redirect('upload')

            # --- Generation via Gemini AI ---
            mcqs = []
            normal_qs = []
            
            if gen_type in ['mcq', 'both']:
                mcqs = generate_mcqs(extracted_text, num_questions=num_questions, difficulty=difficulty)
            
            if gen_type in ['question', 'both']:
                normal_qs = generate_normal_questions(extracted_text, num_questions=num_questions, include_answers=include_answers, difficulty=difficulty)
            
            if not mcqs and not normal_qs:
                if instance and instance.id:
                    instance.delete()
                messages.error(request, "Failed to generate questions from the content. Please try again.")
                return redirect('upload')

            # Create MCQs in database
            if mcqs:
                for q_data in mcqs:
                    Question.objects.create(
                        file=instance,
                        question_text=q_data.get('question', ''),
                        option_a=q_data.get('options', {}).get('A', ''),
                        option_b=q_data.get('options', {}).get('B', ''),
                        option_c=q_data.get('options', {}).get('C', ''),
                        option_d=q_data.get('options', {}).get('D', ''),
                        correct_answer=q_data.get('correct', 'A').upper(),
                        explanation=q_data.get('explanation', '')
                    )

            # Create Normal Questions in database
            if normal_qs:
                for q_data in normal_qs:
                    NormalQuestion.objects.create(
                        file=instance,
                        question_text=q_data.get('question', ''),
                        answer=q_data.get('answer', '')
                    )

            total_generated = len(mcqs) + len(normal_qs)
            messages.success(request, f"Successfully generated {total_generated} question(s) from your content!")
            
            # Redirect based on user's 'After Generation' choice
            if after_generate == 'quiz' and mcqs:
                return redirect('quiz', file_id=instance.id)
            return redirect('review_content', file_id=instance.id)

        except Exception as e:
            # Clean up database entry if error occurs
            if instance and instance.id:
                instance.delete()
            messages.error(request, f"An error occurred while processing: {str(e)}")
            return redirect('upload')
            
    else:
        form = UploadedFileForm()
        
    return render(request, 'user_panel/upload.html', {'form': form})


@login_required
def quiz_view(request, file_id):
    """
    Displays the quiz page containing generated questions for an UploadedFile.
    Configured with a JavaScript countdown timer.
    """
    file_obj = get_object_or_404(UploadedFile, id=file_id, user=request.user)
    questions = file_obj.questions.all().order_by('?')
    
    if not questions.exists():
        messages.error(request, "No questions found for this file. Please regenerate or re-upload.")
        return redirect('dashboard')

    # Assign quiz timer based on the number of questions (e.g. 60 seconds per question)
    timer_seconds = questions.count() * 60

    context = {
        'file_obj': file_obj,
        'questions': questions,
        'timer_seconds': timer_seconds,
    }
    return render(request, 'user_panel/quiz.html', context)


@login_required
def result_view(request, result_id=None):
    """
    Evaluates quiz submission and renders final performance report with detailed correct/wrong answers.
    Can also render a past result if result_id is provided.
    """
    if result_id:
        # Load an existing result from history
        if request.user.is_staff:
            result = get_object_or_404(QuizResult, id=result_id)
        else:
            result = get_object_or_404(QuizResult, id=result_id, user=request.user)
        questions = result.file.questions.all() if result.file else []
        selected_answers = json.loads(result.selected_answers_json) if result.selected_answers_json else {}
    else:
        # Evaluate a new quiz submission
        if request.method != 'POST':
            return redirect('dashboard')
            
        file_id = request.POST.get('file_id')
        file_obj = get_object_or_404(UploadedFile, id=file_id, user=request.user)
        questions = file_obj.questions.all()
        
        score_float = 0.0
        total_questions = questions.count()
        selected_answers = {}
        
        # Grade each question
        for q in questions:
            user_choice = request.POST.get(f"question_{q.id}", "").strip().upper()
            selected_answers[str(q.id)] = user_choice
            
            if user_choice == q.correct_answer:
                score_float += 1.0
            elif user_choice:
                score_float -= 0.25

        if score_float < 0:
            score_float = 0.0

        score = int(round(score_float))
        percentage = round((score_float / total_questions) * 100, 1) if total_questions > 0 else 0
        
        # Save result to database
        result = QuizResult.objects.create(
            user=request.user,
            file=file_obj,
            score=score,
            total_questions=total_questions,
            percentage=percentage,
            selected_answers_json=json.dumps(selected_answers)
        )
        
        messages.success(request, f"Quiz completed! You scored {score}/{total_questions}.")

    # Format feedback message
    if result.percentage >= 80:
        feedback_title = "Excellent Performance!"
        feedback_msg = "Outstanding! You have demonstrating masterly comprehension of the document."
        badge_class = "bg-success"
    elif result.percentage >= 60:
        feedback_title = "Good Work!"
        feedback_msg = "Great effort! A solid comprehension, although there are a few gaps to review."
        badge_class = "bg-info"
    elif result.percentage >= 40:
        feedback_title = "Need Practice!"
        feedback_msg = "Fair. We recommend reviewing the document explanations below to strengthen key facts."
        badge_class = "bg-warning text-dark"
    else:
        feedback_title = "Keep Learning!"
        feedback_msg = "Review the material and re-take the quiz to improve your results!"
        badge_class = "bg-danger"

    # Match each question with selected answers for display
    question_details = []
    for q in questions:
        ans = selected_answers.get(str(q.id), "")
        is_correct = (ans == q.correct_answer)
        
        # Get option texts
        selected_option_text = getattr(q, f"option_{ans.lower()}", "No answer selected") if ans else "No answer selected"
        correct_option_text = getattr(q, f"option_{q.correct_answer.lower()}", "")
        
        question_details.append({
            'question': q,
            'selected': ans,
            'correct': q.correct_answer,
            'is_correct': is_correct,
            'selected_text': selected_option_text,
            'correct_text': correct_option_text,
        })

    context = {
        'result': result,
        'question_details': question_details,
        'feedback_title': feedback_title,
        'feedback_msg': feedback_msg,
        'badge_class': badge_class,
    }
    return render(request, 'user_panel/result.html', context)


@login_required
def review_content_view(request, file_id):
    """
    Displays the generated content (MCQs and Normal Questions) with Copy/Download options.
    """
    if request.user.is_staff:
        file_obj = get_object_or_404(UploadedFile, id=file_id)
    else:
        file_obj = get_object_or_404(UploadedFile, id=file_id, user=request.user)
    mcqs = file_obj.questions.all()
    normal_qs = file_obj.normal_questions.all()

    context = {
        'file_obj': file_obj,
        'mcqs': mcqs,
        'normal_qs': normal_qs,
    }
    return render(request, 'user_panel/review_content.html', context)

@login_required
def file_preview_view(request, file_id):
    """
    Renders a PDF preview page for an uploaded file. Only accessible by owner.
    """
    if request.user.is_staff:
        file_obj = get_object_or_404(UploadedFile, id=file_id)
    else:
        file_obj = get_object_or_404(UploadedFile, id=file_id, user=request.user)
    mcqs = file_obj.questions.all()
    normal_qs = file_obj.normal_questions.all()
    context = {
        'file_obj': file_obj,
        'mcqs': mcqs,
        'normal_qs': normal_qs,
    }
    return render(request, 'user_panel/file_preview.html', context)

@login_required
def history_view(request):
    """
    Renders user's complete history of previous quiz attempts AND all uploaded/generated files.
    """
    results = QuizResult.objects.filter(user=request.user).order_by('-date_taken')
    uploaded_files = UploadedFile.objects.filter(user=request.user).order_by('-uploaded_at')
    return render(request, 'user_panel/history.html', {
        'results': results,
        'uploaded_files': uploaded_files,
    })


# --- Custom Administrative View Controllers ---

@login_required
def admin_dashboard_view(request):
    """
    Renders administrative metrics home page.
    """
    if not request.user.is_staff:
        messages.error(request, "Access denied. Administrative privileges required.")
        return redirect('dashboard')

    # Overall Platform statistics
    total_users = User.objects.count()
    total_files = UploadedFile.objects.count()
    total_questions = Question.objects.count()
    total_attempts = QuizResult.objects.count()

    # Average score across all users
    avg_accuracy = QuizResult.objects.aggregate(Avg('percentage'))['percentage__avg'] or 0.0
    avg_accuracy = round(avg_accuracy, 1)

    # Recent items lists
    recent_users = User.objects.order_by('-date_joined')[:5]
    recent_files = UploadedFile.objects.order_by('-uploaded_at')[:5]
    recent_results = QuizResult.objects.order_by('-date_taken')[:5]

    context = {
        'total_users': total_users,
        'total_files': total_files,
        'total_questions': total_questions,
        'total_attempts': total_attempts,
        'avg_accuracy': avg_accuracy,
        'recent_users': recent_users,
        'recent_files': recent_files,
        'recent_results': recent_results,
    }
    return render(request, 'admin_panel/admin_dashboard.html', context)


@login_required
def admin_manage_users_view(request):
    """
    Renders list of all users and controls.
    """
    if not request.user.is_staff:
        messages.error(request, "Access denied. Administrative privileges required.")
        return redirect('dashboard')

    # Only show normal users, not admins
    users = User.objects.filter(is_staff=False, is_superuser=False).order_by('-date_joined')
    return render(request, 'admin_panel/manage_users.html', {'users': users})


@login_required
def admin_user_detail_view(request, user_id):
    """
    Renders detailed profile, metrics, uploads, and scores for a specific user.
    """
    if not request.user.is_staff:
        messages.error(request, "Access denied. Administrative privileges required.")
        return redirect('dashboard')

    target_user = get_object_or_404(User, id=user_id)
    
    # Calculate user statistics
    uploaded_files = target_user.uploaded_files.all().order_by('-uploaded_at')
    quiz_results = target_user.quiz_results.all().order_by('-date_taken')
    
    # Total questions generated by this user
    mcq_count = sum(f.questions.count() for f in uploaded_files)
    subjective_count = sum(f.normal_questions.count() for f in uploaded_files)
    total_questions = mcq_count + subjective_count
    
    # User's active/suspended/unverified status details
    if target_user.is_active:
        status_label = "Active"
    elif hasattr(target_user, 'otp_verification'):
        status_label = "Unverified"
    else:
        status_label = "Suspended"

    context = {
        'target_user': target_user,
        'uploaded_files': uploaded_files,
        'quiz_results': quiz_results,
        'mcq_count': mcq_count,
        'subjective_count': subjective_count,
        'total_questions': total_questions,
        'status_label': status_label,
    }
    return render(request, 'admin_panel/user_detail.html', context)


@login_required
def admin_toggle_user_view(request, user_id):
    """
    Toggles a user's active status (suspend/activate).
    """
    if not request.user.is_staff:
        messages.error(request, "Access denied. Administrative privileges required.")
        return redirect('dashboard')

    target_user = get_object_or_404(User, id=user_id)
    if target_user == request.user:
        messages.error(request, "You cannot suspend your own administrative account!")
        return redirect('admin_manage_users')

    target_user.is_active = not target_user.is_active
    target_user.save()
    
    status = "activated" if target_user.is_active else "suspended"
    messages.success(request, f"Successfully {status} user account: {target_user.username}")
    return redirect('admin_manage_users')


@login_required
def admin_delete_user_view(request, user_id):
    """
    Deletes a user account from the system (cascades delete related files, questions, results).
    """
    if not request.user.is_staff:
        messages.error(request, "Access denied. Administrative privileges required.")
        return redirect('dashboard')

    target_user = get_object_or_404(User, id=user_id)
    if target_user == request.user:
        messages.error(request, "You cannot delete your own administrative account!")
        return redirect('admin_manage_users')

    username = target_user.username
    target_user.delete()
    messages.success(request, f"Successfully deleted user account: {username}")
    return redirect('admin_manage_users')


@login_required
def admin_manage_files_view(request):
    """
    Renders list of all uploaded documents.
    """
    if not request.user.is_staff:
        messages.error(request, "Access denied. Administrative privileges required.")
        return redirect('dashboard')

    files = UploadedFile.objects.all().order_by('-uploaded_at')
    return render(request, 'admin_panel/manage_files.html', {'files': files})


@login_required
def admin_delete_file_view(request, file_id):
    """
    Deletes an uploaded file (cascades delete related questions & results).
    """
    if not request.user.is_staff:
        messages.error(request, "Access denied. Administrative privileges required.")
        return redirect('dashboard')

    target_file = get_object_or_404(UploadedFile, id=file_id)
    title = target_file.title
    
    # Delete file from storage disk if exists
    if target_file.file and os.path.exists(target_file.file.path):
        try:
            os.remove(target_file.file.path)
        except Exception:
            pass

    target_file.delete()
    messages.success(request, f"Successfully deleted document: '{title}' and all related questions.")
    return redirect('admin_manage_files')


@login_required
def admin_manage_questions_view(request):
    """
    Renders list of all generated questions in database.
    """
    if not request.user.is_staff:
        messages.error(request, "Access denied. Administrative privileges required.")
        return redirect('dashboard')

    questions = Question.objects.all().order_by('file__title')
    normal_questions = NormalQuestion.objects.all().order_by('file__title')
    context = {
        'questions': questions,
        'normal_questions': normal_questions,
    }
    return render(request, 'admin_panel/manage_questions.html', context)


@login_required
def admin_delete_question_view(request, question_id):
    """
    Deletes a specific generated MCQ question.
    """
    if not request.user.is_staff:
        messages.error(request, "Access denied. Administrative privileges required.")
        return redirect('dashboard')

    question = get_object_or_404(Question, id=question_id)
    question.delete()
    messages.success(request, "Successfully deleted question.")
    return redirect('admin_manage_questions')


@login_required
def admin_delete_normal_question_view(request, question_id):
    """
    Deletes a specific generated subjective question.
    """
    if not request.user.is_staff:
        messages.error(request, "Access denied. Administrative privileges required.")
        return redirect('dashboard')

    question = get_object_or_404(NormalQuestion, id=question_id)
    question.delete()
    messages.success(request, "Successfully deleted subjective question.")
    return redirect('admin_manage_questions')


@login_required
def admin_manage_results_view(request):
    """
    Renders list of all previous user test scores.
    """
    if not request.user.is_staff:
        messages.error(request, "Access denied. Administrative privileges required.")
        return redirect('dashboard')

    results = QuizResult.objects.all().order_by('-date_taken')
    return render(request, 'admin_panel/manage_results.html', {'results': results})


@login_required
def admin_delete_result_view(request, result_id):
    """
    Deletes a specific quiz attempt result history.
    """
    if not request.user.is_staff:
        messages.error(request, "Access denied. Administrative privileges required.")
        return redirect('dashboard')

    result = get_object_or_404(QuizResult, id=result_id)
    result.delete()
    messages.success(request, "Successfully deleted quiz result attempt history.")
    return redirect('admin_manage_results')


@login_required
def profile_view(request):
    """
    Renders user profile settings page.
    Allows updating details, password change (requires old password check),
    deleting profile photos, and choosing visual themes.
    """
    profile, created = UserProfile.objects.get_or_create(user=request.user)
    
    # Initialize default forms for both GET and fallback POST contexts
    user_form = UserUpdateForm(instance=request.user)
    profile_form = UserProfileUpdateForm(instance=profile)
    
    if request.method == 'POST':
        action = request.POST.get('action')
        
        # Action 1: Update Details (Name, Email, Avatar)
        if action == 'update_details':
            user_form = UserUpdateForm(request.POST, instance=request.user)
            if user_form.is_valid():
                user_form.save()
                if 'avatar' in request.FILES:
                    profile.avatar = request.FILES['avatar']
                    profile.save()
                messages.success(request, "Account details updated successfully!")
                return redirect('profile')
            else:
                error_details = []
                for field, errs in user_form.errors.items():
                    error_details.append(f"{field}: {', '.join(errs)}")
                messages.error(request, f"Failed to update profile details. Errors: {'; '.join(error_details)}")
                
        # Action 1.5: Update Visual Theme Preferences (Theme, Theme Mode)
        elif action == 'update_theme':
            theme = request.POST.get('theme')
            theme_mode = request.POST.get('theme_mode')
            
            valid_themes = [choice[0] for choice in UserProfile.THEME_CHOICES]
            valid_modes = [choice[0] for choice in UserProfile.THEME_MODE_CHOICES]
            
            if theme in valid_themes and theme_mode in valid_modes:
                profile.theme = theme
                profile.theme_mode = theme_mode
                profile.save()
                messages.success(request, "Theme preferences updated successfully!")
                return redirect('profile')
            else:
                messages.error(request, "Failed to update theme preferences. Invalid choices.")
                
        # Action 1.8: Change Email Request
        elif action == 'change_email_request':
            new_email = request.POST.get('new_email', '').strip()
            
            otp_code = str(secrets.randbelow(900000) + 100000)
            expires_at = timezone.now() + timezone.timedelta(minutes=10)
            request.session['pending_new_email'] = new_email
            OTPVerification.objects.update_or_create(
                user=request.user,
                defaults={'otp_code': otp_code, 'expires_at': expires_at}
            )
            subject = "Verify Your New Email Address - QGenie.AI"
            message = f"Hello {request.user.username},\n\nYou requested to change your email address.\n\nYour 6-digit OTP verification code is: {otp_code}\n\nThis code will expire in 10 minutes.\n\nBest regards,\nQGenie.AI Team"
            try:
                send_mail(subject, message, None, [new_email], fail_silently=False)
                messages.success(request, f"A 6-digit OTP has been sent to {new_email}.")
                return redirect('verify_email_change')
            except Exception as e:
                messages.error(request, f"Failed to send email: {str(e)}")

                
        # Action 2: Change Password
        elif action == 'change_password':
            current_password = request.POST.get('current_password')
            new_password = request.POST.get('new_password')
            confirm_password = request.POST.get('confirm_password')
            
            if not request.user.check_password(current_password):
                messages.error(request, "Incorrect current password. Please try again.")
            elif new_password != confirm_password:
                messages.error(request, "New passwords do not match. Please verify.")
            elif len(new_password) < 6:
                messages.error(request, "Password must be at least 6 characters long.")
            else:
                request.user.set_password(new_password)
                request.user.save()
                login(request, request.user)  # Keep user logged in
                messages.success(request, "Password updated successfully!")
                return redirect('profile')
                
        # Action 2.5: Forgot Current Password Request
        elif action == 'forgot_current_password':
            otp_code = str(secrets.randbelow(900000) + 100000)
            expires_at = timezone.now() + timezone.timedelta(minutes=10)
            OTPVerification.objects.update_or_create(
                user=request.user,
                defaults={'otp_code': otp_code, 'expires_at': expires_at}
            )
            subject = "Reset Your Account Password - OTP Code"
            message = f"Hello {request.user.username},\n\nYou requested to reset your password from your profile.\n\nYour 6-digit OTP reset code is: {otp_code}\n\nThis code will expire in 10 minutes.\n\nBest regards,\nQGenie.AI Team"
            try:
                send_mail(subject, message, None, [request.user.email], fail_silently=False)
                messages.success(request, "An OTP has been sent to your email to reset your password.")
                return redirect('profile_verify_password_reset')
            except Exception as e:
                messages.error(request, f"Failed to send email: {str(e)}")
                
        # Action 3: Delete Avatar
        elif action == 'delete_avatar':
            if profile.avatar:
                # Delete the physical file
                if os.path.exists(profile.avatar.path):
                    os.remove(profile.avatar.path)
                profile.avatar = None
                profile.save()
                messages.success(request, "Profile photo deleted successfully.")
            else:
                messages.info(request, "No profile photo to delete.")
            return redirect('profile')
    return render(request, 'user_panel/profile.html', {
        'user_form': user_form,
        'profile_form': profile_form,
        'profile': profile,
    })


@login_required
def verify_email_change_view(request):
    """
    Verifies the OTP sent to the new email address.
    """
    new_email = request.session.get('pending_new_email')
    if not new_email:
        messages.error(request, "No pending email change request found.")
        return redirect('profile')

    if request.method == 'POST':
        otp_code = "".join(request.POST.get('otp_code', '').split())
        try:
            verification = request.user.otp_verification
            if verification.otp_code == otp_code:
                if verification.is_expired():
                    messages.error(request, "This OTP has expired. Please request a new one.")
                else:
                    # Successful verification! Update email
                    request.user.email = new_email
                    request.user.save()
                    
                    # Delete OTP record
                    verification.delete()
                    del request.session['pending_new_email']
                    
                    messages.success(request, "Your email address has been successfully updated!")
                    return redirect('profile')
            else:
                messages.error(request, "Invalid OTP code. Please try again.")
        except OTPVerification.DoesNotExist:
            messages.error(request, "No OTP request found. Please start over.")
            return redirect('profile')
            
    return render(request, 'user_panel/verify_email_change.html', {'new_email': new_email})


@login_required
def profile_verify_password_reset_view(request):
    """
    Verifies the OTP sent to reset the user's password while they are logged in.
    """
    if request.method == 'POST':
        otp_code = "".join(request.POST.get('otp_code', '').split())
        new_password = request.POST.get('new_password')
        confirm_password = request.POST.get('confirm_password')
        
        if new_password != confirm_password:
            messages.error(request, "Passwords do not match.")
            return render(request, 'user_panel/profile_verify_password_reset.html')
            
        if len(new_password) < 6:
            messages.error(request, "Password must be at least 6 characters.")
            return render(request, 'user_panel/profile_verify_password_reset.html')
            
        try:
            verification = request.user.otp_verification
            if verification.otp_code == otp_code:
                if verification.is_expired():
                    messages.error(request, "This OTP has expired. Please request a new one.")
                else:
                    # Successful verification! Reset password
                    request.user.set_password(new_password)
                    request.user.save()
                    
                    # Keep the user logged in after password change
                    update_session_auth_hash(request, request.user)
                    
                    # Delete OTP record
                    verification.delete()
                    
                    messages.success(request, "Your password has been successfully reset!")
                    return redirect('profile')
            else:
                messages.error(request, "Invalid OTP code. Please try again.")
        except OTPVerification.DoesNotExist:
            messages.error(request, "No OTP request found. Please start over.")
            return redirect('profile')
            
    return render(request, 'user_panel/profile_verify_password_reset.html')





def forgot_password_view(request):
    """
    Renders Forgot Password input page. Generates OTP and emails it.
    """
    if request.user.is_authenticated:
        return redirect('dashboard')
        
    if request.method == 'POST':
        identifier = request.POST.get('identifier', '').strip()
        user = None
        
        try:
            if '@' in identifier:
                users = User.objects.filter(email=identifier)
                if not users.exists():
                    raise User.DoesNotExist
                if users.count() > 1:
                    messages.error(request, "Multiple accounts are registered with this email address. Please use your exact Username to reset your password.")
                    return redirect('forgot_password')
                user = users.first()
            else:
                user = User.objects.get(username=identifier)
        except User.DoesNotExist:
            messages.error(request, "No account found with that username or email address.")
            return redirect('forgot_password')
            
        # Generate 6-digit OTP
        otp_code = str(secrets.randbelow(900000) + 100000)
        expires_at = timezone.now() + timezone.timedelta(minutes=10)
        
        # Save OTP to database
        OTPVerification.objects.update_or_create(
            user=user,
            defaults={'otp_code': otp_code, 'expires_at': expires_at}
        )
        
        # Send Email
        subject = "Reset Your QGenie.AI Password - OTP Code"
        message = f"Hello {user.username},\n\nWe received a request to reset your password.\n\nYour 6-digit OTP reset code is: {otp_code}\n\nThis code will expire in 10 minutes.\n\nIf you did not request this, please ignore this email.\n\nBest regards,\nQGenie.AI Team"
        
        try:
            send_mail(
                subject,
                message,
                None,
                [user.email],
                fail_silently=False,
            )
            messages.success(request, "A password reset OTP has been sent to your registered email.")
        except Exception as e:
            messages.warning(request, f"Failed to send email: {str(e)}. Please contact support.")
            
        return redirect('forgot_password_verify', username=user.username)
        
    return render(request, 'user_panel/forgot_password.html')


def forgot_password_verify_view(request, username):
    """
    Verifies OTP and resets user's password.
    """
    user = get_object_or_404(User, username=username)
    
    if request.method == 'POST':
        otp_code = "".join(request.POST.get('otp_code', '').split())
        new_password = request.POST.get('new_password')
        confirm_password = request.POST.get('confirm_password')
        
        if new_password != confirm_password:
            messages.error(request, "Passwords do not match.")
            return render(request, 'user_panel/forgot_password_verify.html', {'username': username})
            
        if len(new_password) < 6:
            messages.error(request, "Password must be at least 6 characters.")
            return render(request, 'user_panel/forgot_password_verify.html', {'username': username})
            
        try:
            verification = user.otp_verification
            if verification.otp_code == otp_code:
                if verification.is_expired():
                    messages.error(request, "This OTP has expired. Please request a new one.")
                else:
                    # Successful verification! Reset password
                    user.set_password(new_password)
                    user.is_active = True  # Ensure active if they were locked out
                    user.save()
                    
                    # Delete OTP record
                    verification.delete()
                    
                    messages.success(request, "Your password has been successfully reset! You can now log in.")
                    return redirect('login')
            else:
                messages.error(request, "Invalid OTP code. Please try again.")
        except OTPVerification.DoesNotExist:
            messages.error(request, "No OTP request found. Please start over.")
            return redirect('forgot_password')
            
    return render(request, 'user_panel/forgot_password_verify.html', {'username': username})


from django.http import JsonResponse, HttpResponse
from django.views.decorators.csrf import csrf_exempt

@login_required
def chat_with_pdf_view(request, file_id):
    """
    Renders the chat assistant interface and processes user queries against PDF content using AJAX.
    """
    if request.user.is_staff:
        file_obj = get_object_or_404(UploadedFile, id=file_id)
    else:
        file_obj = get_object_or_404(UploadedFile, id=file_id, user=request.user)
    
    # Read text from file
    extracted_text = ""
    try:
        if file_obj.file_type == 'PDF':
            extracted_text = extract_text_from_pdf(file_obj.file.path)
        else:
            with open(file_obj.file.path, 'r', encoding='utf-8', errors='ignore') as f:
                extracted_text = f.read()
    except Exception as e:
        extracted_text = f"Error loading content: {str(e)}"

    if request.method == 'POST':
        # Retrieve message from request
        data = json.loads(request.body)
        user_message = data.get('message', '').strip()
        
        if not user_message:
            return JsonResponse({'error': 'Message cannot be empty.'}, status=400)
            
        # Get chat history from session (specific to this file)
        session_key = f"chat_history_{file_id}"
        chat_history = request.session.get(session_key, [])
        
        # Get AI response
        ai_response = chat_with_document(extracted_text, user_message, chat_history)
        
        # Update and save chat history in session
        chat_history.append({'is_user': True, 'text': user_message})
        chat_history.append({'is_user': False, 'text': ai_response})
        request.session[session_key] = chat_history
        
        return JsonResponse({'answer': ai_response})

    # Render template
    context = {
        'file_obj': file_obj,
    }
    return render(request, 'user_panel/chat_with_pdf.html', context)


@login_required
def summary_view(request, file_id):
    """
    Generates and displays a premium AI study sheet / summary for the uploaded file.
    """
    if request.user.is_staff:
        file_obj = get_object_or_404(UploadedFile, id=file_id)
    else:
        file_obj = get_object_or_404(UploadedFile, id=file_id, user=request.user)
    
    # Check if summary is already cached in database
    if not file_obj.summary_markdown:
        extracted_text = ""
        try:
            if file_obj.file_type == 'PDF':
                extracted_text = extract_text_from_pdf(file_obj.file.path)
            else:
                with open(file_obj.file.path, 'r', encoding='utf-8', errors='ignore') as f:
                    extracted_text = f.read()
            
            # Generate summary via AI
            summary = generate_summary(extracted_text)
            file_obj.summary_markdown = summary
            file_obj.save()
        except Exception as e:
            file_obj.summary_markdown = f"## 📝 Quick Summary\nFailed to generate summary: {str(e)}"
            file_obj.save()

    context = {
        'file_obj': file_obj,
    }
    return render(request, 'user_panel/summary.html', context)



@login_required
def download_content_view(request, file_id, format_type):
    """
    Exports generated MCQs and normal questions as a TXT or MS-Word document.
    """
    if request.user.is_staff:
        file_obj = get_object_or_404(UploadedFile, id=file_id)
    else:
        file_obj = get_object_or_404(UploadedFile, id=file_id, user=request.user)
    questions = file_obj.questions.all()
    normal_questions = file_obj.normal_questions.all()
    
    filename = f"{file_obj.title.replace(' ', '_')}_Generated_Questions"
    
    if format_type == 'txt':
        content = f"=========================================\n"
        content += f"   QGENIE.AI - GENERATED STUDY GUIDE     \n"
        content += f"=========================================\n"
        content += f"Title: {file_obj.title}\n"
        content += f"Date Generated: {file_obj.uploaded_at.strftime('%d %b %Y')}\n\n"
        
        if questions.exists():
            content += f"-----------------------------------------\n"
            content += f"  PART 1: MULTIPLE CHOICE QUESTIONS (MCQs)\n"
            content += f"-----------------------------------------\n\n"
            for i, q in enumerate(questions, 1):
                content += f"Q{i}. {q.question_text}\n"
                content += f"   A. {q.option_a}\n"
                content += f"   B. {q.option_b}\n"
                content += f"   C. {q.option_c}\n"
                content += f"   D. {q.option_d}\n"
                content += f"   Correct Answer: {q.correct_answer}\n"
                if q.explanation:
                    content += f"   Explanation: {q.explanation}\n"
                content += f"\n"
                
        if normal_questions.exists():
            content += f"-----------------------------------------\n"
            content += f"  PART 2: SUBJECTIVE QUESTIONS & ANSWERS \n"
            content += f"-----------------------------------------\n\n"
            for i, q in enumerate(normal_questions, 1):
                content += f"Q{i}. {q.question_text}\n"
                if q.answer:
                    content += f"   Answer: {q.answer}\n"
                content += f"\n"
                
        response = HttpResponse(content, content_type='text/plain; charset=utf-8')
        response['Content-Disposition'] = f'attachment; filename="{filename}.txt"'
        return response
        
    elif format_type == 'doc':
        import io
        from docx import Document
        from docx.shared import Pt, Inches, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        doc = Document()
        
        # Configure standard page margins (1 inch)
        for section in doc.sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)

        # Set default font configuration
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Arial'
        font.size = Pt(11)
        font.color.rgb = RGBColor(0x33, 0x33, 0x33) # Charcoal body text

        # Main Header Title
        title_p = doc.add_paragraph()
        title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title_p.add_run("QGenie.AI - Study Guide")
        title_run.font.name = 'Arial'
        title_run.font.size = Pt(22)
        title_run.font.bold = True
        title_run.font.color.rgb = RGBColor(0x4F, 0x46, 0xE5) # Indigo Accent

        # Document Meta Information
        meta_p = doc.add_paragraph()
        meta_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        meta_run = meta_p.add_run(f"Source Document: {file_obj.title}\nDate Generated: {file_obj.uploaded_at.strftime('%B %d, %Y')}")
        meta_run.font.size = Pt(10)
        meta_run.font.italic = True
        meta_run.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)

        # Add vertical spacing
        spacer_p = doc.add_paragraph()
        spacer_p.paragraph_format.space_after = Pt(20)

        # Part 1: Multiple Choice Questions (MCQs)
        if questions.exists():
            h1 = doc.add_paragraph()
            h1_run = h1.add_run("Part 1: Multiple Choice Questions (MCQs)")
            h1_run.font.size = Pt(16)
            h1_run.font.bold = True
            h1_run.font.color.rgb = RGBColor(0x7C, 0x3A, 0xED) # Purple Accent
            h1.paragraph_format.space_after = Pt(12)

            for i, q in enumerate(questions, 1):
                qp = doc.add_paragraph()
                qp.paragraph_format.space_before = Pt(12)
                qp.paragraph_format.space_after = Pt(6)
                q_run = qp.add_run(f"Q{i}. {q.question_text}")
                q_run.font.bold = True
                q_run.font.size = Pt(11.5)
                q_run.font.color.rgb = RGBColor(0x11, 0x18, 0x27) # Dark Gray / Near Black

                # Render Options A, B, C, D
                options = [
                    ('A', q.option_a),
                    ('B', q.option_b),
                    ('C', q.option_c),
                    ('D', q.option_d)
                ]
                for letter, opt_val in options:
                    if opt_val:
                        opt_p = doc.add_paragraph(style='List Bullet')
                        opt_p.paragraph_format.left_indent = Inches(0.5)
                        opt_p.paragraph_format.space_after = Pt(3)
                        opt_run = opt_p.add_run(f"{letter}. {opt_val}")
                        
                        # Style correct option distinctly in Green if answers should be included
                        if file_obj.include_answers and q.correct_answer == letter:
                            opt_run.font.bold = True
                            opt_run.font.color.rgb = RGBColor(0x05, 0x96, 0x69) # Green Emerald Accent

                # Correct Answer & Explanation
                if file_obj.include_answers:
                    ans_p = doc.add_paragraph()
                    ans_p.paragraph_format.left_indent = Inches(0.5)
                    ans_p.paragraph_format.space_before = Pt(4)
                    ans_p.paragraph_format.space_after = Pt(4)
                    ans_run = ans_p.add_run(f"Correct Option: {q.correct_answer}")
                    ans_run.font.bold = True
                    ans_run.font.color.rgb = RGBColor(0x05, 0x96, 0x69)

                    if q.explanation:
                        exp_p = doc.add_paragraph()
                        exp_p.paragraph_format.left_indent = Inches(0.5)
                        exp_p.paragraph_format.space_after = Pt(12)
                        exp_run_label = exp_p.add_run("Explanation: ")
                        exp_run_label.font.bold = True
                        exp_run_label.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)
                        exp_run_val = exp_p.add_run(q.explanation)
                        exp_run_val.font.italic = True
                        exp_run_val.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)

        # Part 2: Subjective / Normal Questions
        if normal_questions.exists():
            if questions.exists():
                doc.add_page_break()

            h2 = doc.add_paragraph()
            h2_run = h2.add_run("Part 2: Subjective Questions & Answers")
            h2_run.font.size = Pt(16)
            h2_run.font.bold = True
            h2_run.font.color.rgb = RGBColor(0x7C, 0x3A, 0xED) # Purple Accent
            h2.paragraph_format.space_after = Pt(12)

            for i, q in enumerate(normal_questions, 1):
                qp = doc.add_paragraph()
                qp.paragraph_format.space_before = Pt(12)
                qp.paragraph_format.space_after = Pt(6)
                q_run = qp.add_run(f"Q{i}. {q.question_text}")
                q_run.font.bold = True
                q_run.font.size = Pt(11.5)
                q_run.font.color.rgb = RGBColor(0x11, 0x18, 0x27)

                if file_obj.include_answers and q.answer:
                    ans_p = doc.add_paragraph()
                    ans_p.paragraph_format.left_indent = Inches(0.5)
                    ans_p.paragraph_format.space_after = Pt(12)
                    ans_run_label = ans_p.add_run("Answer: ")
                    ans_run_label.font.bold = True
                    ans_run_label.font.color.rgb = RGBColor(0x4F, 0x46, 0xE5)
                    ans_run_val = ans_p.add_run(q.answer)
                    ans_run_val.font.color.rgb = RGBColor(0x37, 0x41, 0x51)

        # Write to byte buffer
        f_stream = io.BytesIO()
        doc.save(f_stream)
        f_stream.seek(0)

        response = HttpResponse(f_stream.getvalue(), content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        response['Content-Disposition'] = f'attachment; filename="{filename}.docx"'
        return response
        
    return redirect('review_content', file_id=file_obj.id)



@login_required
def add_custom_question_view(request, file_id):
    """
    Allows a user to manually add a custom subjective question.
    The AI will automatically answer it based on the document text, and save it as a NormalQuestion.
    """
    if request.user.is_staff:
        file_obj = get_object_or_404(UploadedFile, id=file_id)
    else:
        file_obj = get_object_or_404(UploadedFile, id=file_id, user=request.user)

    if request.method == "POST":
        question_text = request.POST.get('question_text', '').strip()
        if not question_text:
            messages.error(request, "Question text cannot be empty.")
            return redirect('review_content', file_id=file_obj.id)
            
        try:
            # Extract text from file
            extracted_text = ""
            if file_obj.file_type == 'PDF':
                extracted_text = extract_text_from_pdf(file_obj.file.path)
            else:
                with open(file_obj.file.path, 'r', encoding='utf-8', errors='ignore') as f_read:
                    extracted_text = f_read.read()
                    
            # Use AI to answer the specific question based on the document
            ai_answer = chat_with_document(extracted_text, question_text)
            
            # Save it to the database
            NormalQuestion.objects.create(
                file=file_obj,
                question_text=question_text,
                answer=ai_answer
            )
            
            messages.success(request, "Custom question added and answered successfully by QGenie!")
        except Exception as e:
            messages.error(request, f"Failed to generate answer for your question: {str(e)}")
            
    return redirect('review_content', file_id=file_obj.id)
